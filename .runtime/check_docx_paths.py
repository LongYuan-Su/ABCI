from __future__ import annotations

import json
import re
from pathlib import Path

import docx


ROOT = Path(r"C:\Users\SLY13\Desktop\MetaBCI_Integrated_Initial_Version")
DOCX_PATH = ROOT / "MetaBCI初赛_吞咽项目.docx"

KNOWN_TOP = {
    "metabci",
    "applications",
    "demos",
    "docs",
    "logs",
    ".runtime",
    "tools",
    "hardware",
    "models",
    "tests",
    "images",
}

PATH_PAT = re.compile(
    r"(?:(?:[A-Za-z]:)?[A-Za-z0-9_.\-\u4e00-\u9fff<>*]+[\\/])+"
    r"[A-Za-z0-9_.\-\u4e00-\u9fff<>*]+"
)
FILE_PAT = re.compile(
    r"(?<![A-Za-z0-9_\\/.-])"
    r"([A-Za-z0-9_\-\u4e00-\u9fff]+\."
    r"(?:py|md|npy|json|pt|db|mp3|csv|docx|yml|yaml|toml|txt|ini|png|jpg|jpeg|pdf))"
    r"(?![A-Za-z0-9_.-])"
)


def iter_paragraphs(document: docx.Document):
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def clean_token(token: str) -> str:
    return token.strip("` \t\r\n，。；;、：:（）()[]【】<>\"“”'")


def resolve_token(token: str):
    raw = token.replace("/", "\\")
    if "<" in raw or ">" in raw or "*" in raw:
        return None, "template"
    if re.match(r"^[A-Za-z]:\\", raw):
        return Path(raw), "absolute"

    first = raw.split("\\")[0]
    if first in KNOWN_TOP or raw in {
        "README.md",
        "LICENSE",
        "pyproject.toml",
        "environment.yml",
        "requirements.txt",
        "setup.py",
    }:
        return ROOT / raw, "relative"

    # Path suffix such as "gui\\eeg_display.py"; search by suffix.
    if "\\" in raw:
        raw_parts = [part for part in raw.split("\\") if part]
        matches = []
        for path in ROOT.rglob(raw_parts[-1]):
            parts = set(path.parts)
            if "__pycache__" in parts or ".git" in parts:
                continue
            if list(path.parts[-len(raw_parts):]) == raw_parts:
                matches.append(path)
        return (matches[0] if matches else ROOT / raw), "suffix"

    matches = []
    for path in ROOT.rglob(raw):
        parts = set(path.parts)
        if "__pycache__" in parts or ".git" in parts:
            continue
        matches.append(path)
    return (matches[0] if matches else ROOT / raw), "bare"


def is_real_path_candidate(token: str) -> bool:
    raw = token.replace("/", "\\")
    if "\\" not in raw:
        return True
    if re.match(r"^[A-Za-z]:\\", raw):
        return True
    first = raw.split("\\")[0]
    if first in KNOWN_TOP:
        return True
    # Treat slash-separated file lists like data.npy/labels.json/meta.json as
    # separate bare files instead of one nested path.
    parts = [part for part in raw.split("\\") if part]
    if len(parts) > 1 and all("." in part for part in parts):
        return False
    return "." in parts[-1]


def main() -> None:
    document = docx.Document(str(DOCX_PATH))
    seen: dict[str, dict] = {}
    for line_no, paragraph in enumerate(iter_paragraphs(document), 1):
        line = paragraph.text
        for match in PATH_PAT.finditer(line):
            token = clean_token(match.group(0))
            if not token or not is_real_path_candidate(token):
                continue
            path, kind = resolve_token(token)
            exists = None if kind == "template" else bool(path and path.exists())
            seen.setdefault(
                token,
                {
                    "kind": kind,
                    "exists": exists,
                    "line_samples": [],
                    "resolved": str(path) if path else "",
                },
            )["line_samples"].append(line_no)

        for match in FILE_PAT.finditer(line):
            token = clean_token(match.group(1))
            if not token:
                continue
            if any(token in existing for existing in seen):
                continue
            path, kind = resolve_token(token)
            exists = None if kind == "template" else bool(path and path.exists())
            seen.setdefault(
                token,
                {
                    "kind": kind,
                    "exists": exists,
                    "line_samples": [],
                    "resolved": str(path) if path else "",
                },
            )["line_samples"].append(line_no)

    items = []
    for token, value in sorted(seen.items(), key=lambda item: (str(item[1]["exists"]), item[0].lower())):
        value["mentions"] = len(value["line_samples"])
        value["lines"] = sorted(set(value["line_samples"]))[:10]
        items.append((token, value))

    payload = {
        "count": len(items),
        "missing": [token for token, value in items if value["exists"] is False],
        "templates": [token for token, value in items if value["exists"] is None],
        "items": items,
    }
    (ROOT / ".runtime" / "path_check_report.json").write_text(
        json.dumps(payload, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    print(json.dumps(payload, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
