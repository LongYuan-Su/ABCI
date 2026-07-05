from __future__ import annotations

import shutil
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\Users\SLY13\Desktop\MetaBCI_Integrated_Initial_Version")
DOCX_PATH = ROOT / "MetaBCI初赛_吞咽项目.docx"
BACKUP_PATH = ROOT / "MetaBCI初赛_吞咽项目.before_path_update.docx"


REPLACEMENTS = [
    (
        r"metabci\brainstim、metabci\brainflow、metabci\brainda、tools、hardware、models、demos\swallow_demos目录",
        r"metabci\brainstim、metabci\brainflow、metabci\brainda、applications\swallow_bci目录",
    ),
    (
        r"demos\swallow_demos\swallow_demos\run_all_demos.py",
        r"applications\swallow_bci\demos\swallow_demos\run_all_demos.py",
    ),
    (
        r"demos\swallow_demos\ demo_03_play_paradigm2.py",
        r"applications\swallow_bci\demos\swallow_demos\demo_03_play_paradigm2.py",
    ),
    (
        r"demos\swallow_demos\demo_01_main_gui_demo.py",
        r"applications\swallow_bci\demos\swallow_demos\demo_01_main_gui_demo.py",
    ),
    (
        r"demos\swallow_demos\demo_02_play_paradigm1.py",
        r"applications\swallow_bci\demos\swallow_demos\demo_02_play_paradigm1.py",
    ),
    (
        r"demos\swallow_demos\demo_03_play_paradigm2.py",
        r"applications\swallow_bci\demos\swallow_demos\demo_03_play_paradigm2.py",
    ),
    (
        r"demos\swallow_demos\demo_04_control_evaluation.py",
        r"applications\swallow_bci\demos\swallow_demos\demo_04_control_evaluation.py",
    ),
    (
        r"demos\swallow_demos\run_all_demos.py",
        r"applications\swallow_bci\demos\swallow_demos\run_all_demos.py",
    ),
    (
        r"demos\swallow_demos",
        r"applications\swallow_bci\demos\swallow_demos",
    ),
    (
        "demos/swallow_demos",
        "applications/swallow_bci/demos/swallow_demos",
    ),
    (
        r"hardware\esp32b_controller\src",
        r"applications\swallow_bci\hardware\esp32b_controller\src",
    ),
    (
        "hardware/esp32b_controller",
        "applications/swallow_bci/hardware/esp32b_controller",
    ),
    (
        "models/swallow_classifier",
        "applications/swallow_bci/models/swallow_classifier",
    ),
    (
        r"tools\esp32b_control_gui.py",
        r"applications\swallow_bci\tools\esp32b_control_gui.py",
    ),
    (
        "tools/esp32b_control_gui.py",
        "applications/swallow_bci/tools/esp32b_control_gui.py",
    ),
    (
        r"metabci\brainflow\sources.py",
        r"metabci\brainflow\acquisition\sources.py",
    ),
    (
        r"metabci\brainflow\recorder.py",
        r"metabci\brainflow\acquisition\recorder.py",
    ),
    (
        r"metabci\brainflow\closed_loop.py",
        r"metabci\brainflow\control\closed_loop.py",
    ),
    (
        r"metabci\brainflow\online_swallow_control.py",
        r"metabci\brainflow\control\online_swallow_control.py",
    ),
    (
        r"metabci\brainflow\assessment.py",
        r"metabci\brainflow\processing\assessment.py",
    ),
    (
        r"metabci\brainflow\decoder.py",
        r"metabci\brainflow\processing\decoder.py",
    ),
]

REPLACEMENT_MAP = dict(REPLACEMENTS)
REPLACEMENT_RE = re.compile(
    "|".join(re.escape(old) for old, _new in sorted(REPLACEMENTS, key=lambda pair: len(pair[0]), reverse=True))
)


def iter_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            for paragraph in nested_cell.paragraphs:
                                yield paragraph


def rebuild_paragraph_text(paragraph: Paragraph, text: str) -> None:
    first_r_pr = None
    for run in paragraph.runs:
        if run.text:
            first_r_pr = deepcopy(run._r.rPr) if run._r.rPr is not None else None
            break

    p = paragraph._p
    for run in list(paragraph.runs):
        p.remove(run._r)

    new_r = OxmlElement("w:r")
    if first_r_pr is not None:
        new_r.append(first_r_pr)
    p.append(new_r)
    new_run = paragraph.runs[-1]
    new_run.text = text
    new_run.font.highlight_color = None


def clear_yellow(paragraph: Paragraph) -> int:
    cleared = 0
    for run in paragraph.runs:
        if run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
            run.font.highlight_color = None
            cleared += 1
    return cleared


def update_text(text: str) -> tuple[str, int]:
    changed = 0

    def replace_once(match: re.Match[str]) -> str:
        nonlocal changed
        changed += 1
        return REPLACEMENT_MAP[match.group(0)]

    return REPLACEMENT_RE.sub(replace_once, text), changed


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    doc = Document(str(DOCX_PATH))
    changed_paragraphs = 0
    replacement_count = 0
    yellow_cleared = 0

    for paragraph in iter_paragraphs(doc):
        new_text, count = update_text(paragraph.text)
        if count:
            rebuild_paragraph_text(paragraph, new_text)
            changed_paragraphs += 1
            replacement_count += count
        yellow_cleared += clear_yellow(paragraph)

    doc.save(str(DOCX_PATH))
    print(f"changed_paragraphs={changed_paragraphs}")
    print(f"replacement_count={replacement_count}")
    print(f"yellow_cleared={yellow_cleared}")
    print(f"backup={BACKUP_PATH}")


if __name__ == "__main__":
    main()
