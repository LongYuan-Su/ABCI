from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.text.run import Run


ROOT = Path(r"C:\Users\SLY13\Desktop\MetaBCI_Integrated_Initial_Version")
DOCX_PATH = ROOT / "MetaBCI初赛_吞咽项目.docx"
BACKUP_PATH = ROOT / "MetaBCI初赛_吞咽项目.before_path_highlight.docx"


# These are the vetted references that point to files/directories that no longer
# exist at the paths written in the document. Generic output filenames and
# template paths such as logs\<患者ID>\data.npy are intentionally excluded.
MISSING_REFERENCES = {
    r"demos/swallow_demos",
    r"demos\swallow_demos",
    r"demos\swallow_demos\demo_01_main_gui_demo.py",
    r"demos\swallow_demos\demo_02_play_paradigm1.py",
    r"demos\swallow_demos\demo_03_play_paradigm2.py",
    r"demos\swallow_demos\ demo_03_play_paradigm2.py",
    r"demos\swallow_demos\demo_04_control_evaluation.py",
    r"demos\swallow_demos\run_all_demos.py",
    r"demos\swallow_demos\swallow_demos\run_all_demos.py",
    r"hardware/esp32b_controller",
    r"hardware\esp32b_controller\src",
    r"metabci\brainflow\assessment.py",
    r"metabci\brainflow\closed_loop.py",
    r"metabci\brainflow\decoder.py",
    r"metabci\brainflow\online_swallow_control.py",
    r"metabci\brainflow\recorder.py",
    r"metabci\brainflow\sources.py",
    r"models/swallow_classifier",
    r"tools/esp32b_control_gui.py",
    r"tools\esp32b_control_gui.py",
}


TOKEN_RE = re.compile(
    "|".join(re.escape(token) for token in sorted(MISSING_REFERENCES, key=len, reverse=True))
)


def _insert_run_after(run: Run, text: str, highlighted: bool, base_r_pr) -> Run:
    new_r = OxmlElement("w:r")
    if base_r_pr is not None:
        new_r.append(deepcopy(base_r_pr))
    run._r.addnext(new_r)
    new_run = Run(new_r, run._parent)
    new_run.text = text
    if highlighted:
        new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    else:
        new_run.font.highlight_color = None
    return new_run


def _highlight_run_matches(run: Run) -> int:
    text = run.text
    matches = list(TOKEN_RE.finditer(text))
    if not matches:
        return 0
    base_r_pr = deepcopy(run._r.rPr) if run._r.rPr is not None else None

    pieces: list[tuple[str, bool]] = []
    cursor = 0
    for match in matches:
        if match.start() > cursor:
            pieces.append((text[cursor : match.start()], False))
        pieces.append((match.group(0), True))
        cursor = match.end()
    if cursor < len(text):
        pieces.append((text[cursor:], False))

    first_text, first_highlighted = pieces[0]
    run.text = first_text
    if first_highlighted:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    else:
        run.font.highlight_color = None

    current = run
    for piece_text, highlighted in pieces[1:]:
        if not piece_text:
            continue
        current = _insert_run_after(current, piece_text, highlighted, base_r_pr)
    return len(matches)


def iter_paragraphs(doc: Document):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                for nested_table in cell.tables:
                    for nested_row in nested_table.rows:
                        for nested_cell in nested_row.cells:
                            for paragraph in nested_cell.paragraphs:
                                yield paragraph


def highlight_paragraph(paragraph: Paragraph) -> int:
    count = 0
    for run in list(paragraph.runs):
        count += _highlight_run_matches(run)
    return count


def clear_unrelated_yellow(paragraph: Paragraph) -> int:
    cleared = 0
    for run in paragraph.runs:
        if run.font.highlight_color == WD_COLOR_INDEX.YELLOW and not TOKEN_RE.search(run.text):
            run.font.highlight_color = None
            cleared += 1
    return cleared


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    if not BACKUP_PATH.exists():
        shutil.copy2(DOCX_PATH, BACKUP_PATH)

    doc = Document(str(DOCX_PATH))
    highlighted_count = 0
    paragraphs_with_hits = 0
    for paragraph in iter_paragraphs(doc):
        count = highlight_paragraph(paragraph)
        if count:
            highlighted_count += count
            paragraphs_with_hits += 1
    cleared_unrelated = sum(clear_unrelated_yellow(paragraph) for paragraph in iter_paragraphs(doc))

    doc.save(str(DOCX_PATH))
    print(f"highlighted_occurrences={highlighted_count}")
    print(f"paragraphs_with_hits={paragraphs_with_hits}")
    print(f"cleared_unrelated_yellow_runs={cleared_unrelated}")
    print(f"backup={BACKUP_PATH}")


if __name__ == "__main__":
    main()
